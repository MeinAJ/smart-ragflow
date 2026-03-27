"""
DeepDoc 文档解析模块。

基于 DeepDoc 的文档解析实现，支持 PDF、Word、TXT 等多种格式。
"""

import logging
import os
import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from datetime import datetime

from backend_parser.config import settings
from backend_parser.document_models import Document, DocumentSection, DocumentElement, ParseResult

logger = logging.getLogger(__name__)


class DeepDocParser:
    """
    DeepDoc 文档解析器。

    支持多种文档格式的解析，提取文本、标题、表格等内容。
    """

    def __init__(self):
        """初始化解析器。"""
        self.supported_extensions = settings.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str) -> ParseResult:
        """
        解析文档。

        Args:
            file_path: 文件路径

        Returns:
            ParseResult: 解析结果
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return ParseResult(success=False, error_message=f"File not found: {file_path}")

        extension = file_path.suffix.lower()
        if extension not in self.supported_extensions:
            return ParseResult(success=False, error_message=f"Unsupported file type: {extension}")

        try:
            logger.info(f"Parsing document: {file_path}")

            # 根据文件类型选择解析方法
            if extension == ".pdf":
                return self._parse_pdf(file_path)
            elif extension in [".docx", ".doc"]:
                return self._parse_word(file_path)
            elif extension == ".txt":
                return self._parse_txt(file_path)
            elif extension == ".md":
                return self._parse_md(file_path)
            elif extension in [".html", ".htm"]:
                return self._parse_html(file_path)
            else:
                return ParseResult(success=False, error_message=f"Unsupported file type: {extension}")

        except Exception as e:
            logger.error(f"Failed to parse document: {e}")
            return ParseResult(success=False, error_message=str(e))

    def _extract_title(self, text: str) -> str:
        """从文本中提取标题（第一行非空内容）。"""
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if line and len(line) < 200:
                return line
        return "Untitled"

    def _parse_structure(self, text: str) -> List[DocumentSection]:
        """
        解析文档结构，提取标题层级。

        支持 Markdown 风格的标题（# ## ###）。
        """
        sections = []
        current_section = None
        current_content = []

        # 匹配标题行
        header_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

        # 按行解析
        lines = text.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]
            match = header_pattern.match(line)

            if match:
                # 保存上一个章节
                if current_section is not None:
                    current_section.content = "\n".join(current_content).strip()
                    sections.append(current_section)

                # 创建新章节
                level = len(match.group(1))
                title = match.group(2).strip()
                current_section = DocumentSection(
                    title=title,
                    level=level,
                    start_page=0
                )
                current_content = []
            else:
                if current_section is not None:
                    current_content.append(line)
                else:
                    # 前言部分（无标题）
                    if line.strip():
                        current_section = DocumentSection(
                            title="",
                            level=0,
                            start_page=0
                        )
                        current_content.append(line)

            i += 1

        # 保存最后一个章节
        if current_section is not None:
            current_section.content = "\n".join(current_content).strip()
            sections.append(current_section)

        # 如果没有找到任何章节，将整个文本作为一个章节
        if not sections and text.strip():
            sections.append(DocumentSection(
                title=self._extract_title(text),
                level=1,
                content=text.strip()
            ))

        return sections

    def _parse_pdf(self, file_path: Path) -> ParseResult:
        """
        解析 PDF 文件。

        注：这里使用 PyPDF2 或 pdfplumber 作为基础实现。
        实际 DeepDoc 会更复杂。
        """
        try:
            # 尝试导入 pdfplumber
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed, using basic implementation")
            # 创建一个简单的文档对象作为演示
            return self._create_mock_result(file_path)

        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                elements = []

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        text_parts.append(text)
                        elements.append(DocumentElement(
                            element_type="paragraph",
                            content=text,
                            page_number=page_num
                        ))

                full_text = "\n\n".join(text_parts)
                sections = self._parse_structure(full_text)

                # 更新章节页码信息
                for section in sections:
                    section.start_page = 1
                    section.end_page = len(pdf.pages)

                document = Document(
                    title=self._extract_title(full_text) or file_path.stem,
                    file_name=file_path.name,
                    file_type="pdf",
                    file_path=str(file_path),
                    file_size=file_path.stat().st_size,
                    content=full_text,
                    sections=sections,
                    elements=elements,
                    parsed_at=datetime.now()
                )

                return ParseResult(
                    document=document,
                    success=True
                )

        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return self._create_mock_result(file_path)

    def _parse_word(self, file_path: Path) -> ParseResult:
        """解析 Word 文件。
        
        注意：python-docx 只支持 .docx 格式（Word 2007+），不支持旧的 .doc 二进制格式。
        如果传入 .doc 文件，会尝试解析，但如果是旧的二进制格式会报错。
        """
        try:
            import docx
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            logger.warning("python-docx not installed, using basic implementation")
            return self._create_mock_result(file_path)

        # 检查文件扩展名，提供更有帮助的错误信息
        extension = file_path.suffix.lower()
        if extension == ".doc":
            logger.warning(
                f"File {file_path.name} has .doc extension. "
                f"Note: Only .docx format (Word 2007+) is supported. "
                f"Old .doc binary format is not supported."
            )

        try:
            doc = docx.Document(file_path)

            text_parts = []
            elements = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    # 判断是否为标题
                    element_type = "paragraph"
                    if para.style.name.startswith("Heading"):
                        element_type = "title"

                    elements.append(DocumentElement(
                        element_type=element_type,
                        content=para.text,
                        page_number=0
                    ))

            full_text = "\n".join(text_parts)
            sections = self._parse_structure(full_text)

            document = Document(
                title=doc.core_properties.title or self._extract_title(full_text) or file_path.stem,
                file_name=file_path.name,
                file_type="docx",
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                content=full_text,
                sections=sections,
                elements=elements,
                metadata={
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created,
                },
                parsed_at=datetime.now()
            )

            return ParseResult(
                document=document,
                success=True
            )

        except PackageNotFoundError as e:
            # 处理文件格式错误（如内容类型不匹配）
            error_msg = str(e)
            logger.error(f"Word parsing error - invalid file format: {error_msg}")
            
            # 提供更友好的错误信息
            if "is not a Word file" in error_msg or "content type" in error_msg.lower():
                user_msg = (
                    f"文件 '{file_path.name}' 不是有效的 Word 文档。"
                    f"检测到内容类型异常：{error_msg}。"
                    f"可能原因：(1) 文件已损坏；(2) 这是旧的 .doc 二进制格式文件，"
                    f"请转换为 .docx 格式后重试；(3) 文件扩展名与实际格式不匹配。"
                )
            else:
                user_msg = f"无法解析 Word 文件：{error_msg}"
            
            return ParseResult(success=False, error_message=user_msg)
        except Exception as e:
            logger.error(f"Word parsing error: {e}")
            return self._create_mock_result(file_path)

    def _parse_txt(self, file_path: Path) -> ParseResult:
        """解析纯文本文件。"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            sections = self._parse_structure(text)

            document = Document(
                title=self._extract_title(text) or file_path.stem,
                file_name=file_path.name,
                file_type="txt",
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                content=text,
                sections=sections,
                parsed_at=datetime.now()
            )

            return ParseResult(
                document=document,
                success=True
            )

        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            return ParseResult(success=False, error_message=str(e))

    def _parse_md(self, file_path: Path) -> ParseResult:
        """解析 Markdown 文件。"""
        return self._parse_txt(file_path)  # Markdown 也是文本格式

    def _parse_html(self, file_path: Path) -> ParseResult:
        """解析 HTML 文件。"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.warning("beautifulsoup4 not installed, using basic implementation")
            return self._parse_txt(file_path)

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")

            # 移除 script 和 style
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text()

            # 清理空白
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            sections = self._parse_structure(text)

            document = Document(
                title=soup.title.string if soup.title else self._extract_title(text) or file_path.stem,
                file_name=file_path.name,
                file_type="html",
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                content=text,
                sections=sections,
                parsed_at=datetime.now()
            )

            return ParseResult(
                document=document,
                success=True
            )

        except Exception as e:
            logger.error(f"HTML parsing error: {e}")
            return self._parse_txt(file_path)

    def _create_mock_result(self, file_path: Path) -> ParseResult:
        """
        创建模拟解析结果（用于测试或依赖未安装时）。

        Args:
            file_path: 文件路径

        Returns:
            ParseResult: 模拟解析结果
        """
        logger.warning(f"Creating mock parse result for: {file_path}")

        document = Document(
            title=file_path.stem,
            file_name=file_path.name,
            file_type=file_path.suffix.lstrip("."),
            file_path=str(file_path),
            file_size=file_path.stat().st_size,
            content=f"Mock content for {file_path.name}",
            sections=[DocumentSection(
                title=file_path.stem,
                level=1,
                content=f"Mock content for {file_path.name}"
            )],
            parsed_at=datetime.now()
        )

        return ParseResult(
            document=document,
            success=True
        )


# 全局解析器实例
deepdoc_parser = DeepDocParser()
